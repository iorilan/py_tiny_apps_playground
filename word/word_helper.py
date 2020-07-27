import docx

def read_all(file_name):
    #print(file_name)
    doc = docx.Document(file_name)
    for para in doc.paragraphs:
        print(para.text)
        print('-----')
        for run in para.runs:
            print(run.text)
            print('^^^^^')

"""
page
    header1
    header2
    para1
        run1
        run2
    para2
        run3..
"""
def create_new(pages, file_name):
    mydoc = docx.Document()
    for p in pages:
        headings = p['headings']
        for h in headings:
            mydoc.add_heading(h[0],int(h[1]))
        paras = p['paragraph']
        for p1 in paras:
            para_ = mydoc.add_paragraph(p1['txt'])
            for r in p1['runs']:
                para_.add_run(r)
        mydoc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    mydoc.save(file_name)
